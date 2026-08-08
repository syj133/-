# -*- coding: utf-8 -*-
"""Extract text from downloaded Huashu Cup official PDF/DOCX files."""
import os
import sys

from pypdf import PdfReader
from docx import Document

BASE = os.path.join(os.getcwd(), "优秀论文", "华数杯专项")


def main():
    for name in sorted(os.listdir(BASE)):
        path = os.path.join(BASE, name)
        print("=" * 90)
        print(f"FILE: {name}")
        print("=" * 90)
        if name.lower().endswith(".pdf"):
            reader = PdfReader(path)
            print(f"[PDF, {len(reader.pages)} pages]")
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                print(f"\n----- page {i + 1} -----")
                print(text)
        elif name.lower().endswith(".docx"):
            doc = Document(path)
            print("[DOCX]")
            for p in doc.paragraphs:
                if p.text.strip():
                    print(p.text)
            for i, table in enumerate(doc.tables):
                print(f"\n[table {i + 1}]")
                for row in table.rows:
                    cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                    print(" | ".join(cells))


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    main()
